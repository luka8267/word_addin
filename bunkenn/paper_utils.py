import io
import os
import time

from docx import Document
from postgrest.exceptions import APIError

BUCKET_NAME = "paper-pdfs"
READING_STATUSES = ["未読", "読書中", "読了", "再読したい", "引用予定"]
SORT_OPTIONS = ["追加順", "年（新しい順）", "年（古い順）", "タイトル", "ステータス"]


def normalize_doi(doi):
    return (doi or "").strip()


def normalize_tag_input(tags_text):
    seen = set()
    normalized = []
    for tag in (tags_text or "").split(","):
        value = tag.strip()
        if value and value not in seen:
            seen.add(value)
            normalized.append(value)
    return normalized


def fetch_user_papers(supabase, user_id, columns="*"):
    try:
        return (
            supabase.table("papers")
            .select(columns)
            .eq("user_id", user_id)
            .order("display_order")
            .execute()
        )
    except APIError as error:
        error_text = str(error).lower()
        if "uuid" in error_text or "user_id" in error_text:
            raise RuntimeError(
                "papers.user_id と認証ユーザーの紐づけ、または RLS 設定を確認してください。"
            ) from error
        raise


def search_user_papers(supabase, user_id, keyword, columns="id, title, authors, year"):
    normalized_keyword = (keyword or "").strip()
    query = (
        supabase.table("papers")
        .select(columns)
        .eq("user_id", user_id)
        .order("display_order")
    )

    if normalized_keyword:
        escaped_keyword = normalized_keyword.replace("%", "\\%").replace(",", "\\,")
        query = query.or_(
            f"title.ilike.%{escaped_keyword}%,authors.ilike.%{escaped_keyword}%"
        )

    return query.execute()


def sort_papers_dataframe(df, sort_option):
    if df.empty:
        return df

    sorted_df = df.copy()

    if sort_option == "年（新しい順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=False)
    elif sort_option == "年（古い順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=True)
    elif sort_option == "タイトル":
        sorted_df = sorted_df.sort_values(by="title", ascending=True)
    elif sort_option == "ステータス":
        sorted_df = sorted_df.sort_values(by="status", ascending=True)

    sorted_df = sorted_df.reset_index(drop=True)
    sorted_df["ref_no"] = sorted_df.index + 1
    return sorted_df


def make_word_citation(row, style="APA"):
    authors = row.get("authors", "")
    year = row.get("year", "")
    title = row.get("title", "")
    journal = row.get("journal", "")
    doi = row.get("doi", "")

    if style == "APA":
        citation = f"{authors} ({year}). {title}. {journal}."
        if doi:
            citation += f" https://doi.org/{doi}"
    elif style == "Vancouver":
        citation = f"{authors}. {title}. {journal}. {year}."
        if doi:
            citation += f" doi:{doi}"
    elif style == "Nature":
        citation = f"{authors} {title}. {journal} ({year})."
        if doi:
            citation += f" https://doi.org/{doi}"
    else:
        citation = f"{authors} ({year}). {title}. {journal}."

    return citation


def build_bibliography_entries(papers, style="APA", numbered=True):
    entries = []
    for index, paper in enumerate(papers, start=1):
        citation = make_word_citation(paper, style=style)
        if numbered:
            citation = f"[{index}] {citation}"
        entries.append(citation)
    return entries


def export_to_word_bytes(papers, style="APA", title="参考文献", numbered=True):
    doc = Document()
    doc.add_heading(title, 0)

    for entry in build_bibliography_entries(papers, style=style, numbered=numbered):
        doc.add_paragraph(entry)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def fetch_papers_by_ids(supabase, user_id, paper_ids, columns="*"):
    if not paper_ids:
        return []

    result = (
        supabase.table("papers")
        .select(columns)
        .eq("user_id", user_id)
        .in_("id", paper_ids)
        .execute()
    )

    papers_by_id = {paper["id"]: paper for paper in (result.data or [])}
    return [papers_by_id[paper_id] for paper_id in paper_ids if paper_id in papers_by_id]


def upload_pdf_to_storage(supabase, pdf_file, user_id):
    filename = pdf_file.name
    name, ext = os.path.splitext(filename)
    safe_name = f"{name}_{int(time.time())}{ext}"
    storage_path = f"{user_id}/{safe_name}"

    supabase.storage.from_(BUCKET_NAME).upload(
        path=storage_path,
        file=pdf_file.read(),
        file_options={"content-type": "application/pdf"},
    )
    return storage_path


def create_pdf_signed_url(supabase, storage_path, expires_in=3600):
    if not isinstance(storage_path, str) or not storage_path.strip():
        return None

    response = (
        supabase.storage.from_(BUCKET_NAME).create_signed_url(
            storage_path,
            expires_in,
        )
    )

    if isinstance(response, dict):
        return response.get("signedURL") or response.get("signedUrl")
    return None


def delete_pdf_from_storage(supabase, storage_path):
    if storage_path:
        supabase.storage.from_(BUCKET_NAME).remove([storage_path])


def get_or_create_tag_id(supabase, tag_name):
    tag_result = (
        supabase.table("tags")
        .select("id")
        .eq("name", tag_name)
        .limit(1)
        .execute()
    )

    if tag_result.data:
        return tag_result.data[0]["id"]

    new_tag = supabase.table("tags").insert({"name": tag_name}).execute()
    return new_tag.data[0]["id"]


def save_tags_for_paper(supabase, paper_id, tags_text):
    for tag_name in normalize_tag_input(tags_text):
        tag_id = get_or_create_tag_id(supabase, tag_name)
        supabase.table("paper_tags").upsert(
            {"paper_id": paper_id, "tag_id": tag_id}
        ).execute()


def get_tag_map_for_papers(supabase, paper_ids):
    if not paper_ids:
        return {}

    paper_tag_result = (
        supabase.table("paper_tags")
        .select("paper_id, tag_id")
        .in_("paper_id", paper_ids)
        .execute()
    )

    paper_tags = paper_tag_result.data or []
    if not paper_tags:
        return {}

    tag_ids = sorted({row["tag_id"] for row in paper_tags})
    tag_result = supabase.table("tags").select("id, name").in_("id", tag_ids).execute()
    tag_name_map = {row["id"]: row["name"] for row in (tag_result.data or [])}

    tag_map = {paper_id: [] for paper_id in paper_ids}
    for row in paper_tags:
        tag_name = tag_name_map.get(row["tag_id"])
        if tag_name:
            tag_map.setdefault(row["paper_id"], []).append(tag_name)

    return tag_map


def move_paper(supabase, user_id, paper_id, display_order, direction):
    operator = "lt" if direction == "up" else "gt"
    descending = direction == "up"

    neighbor_result = (
        getattr(
            supabase.table("papers")
            .select("id, display_order")
            .eq("user_id", user_id),
            operator,
        )("display_order", display_order)
        .order("display_order", desc=descending)
        .limit(1)
        .execute()
    )

    if not neighbor_result.data:
        return

    neighbor = neighbor_result.data[0]

    (
        supabase.table("papers")
        .update({"display_order": neighbor["display_order"]})
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )

    (
        supabase.table("papers")
        .update({"display_order": display_order})
        .eq("id", neighbor["id"])
        .eq("user_id", user_id)
        .execute()
    )


def update_paper_details(supabase, user_id, paper_id, status, notes):
    (
        supabase.table("papers")
        .update({"status": status, "notes": notes})
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )


def delete_paper(supabase, user_id, row):
    pdf_path = row.get("pdf_path")
    if isinstance(pdf_path, str) and pdf_path.strip():
        delete_pdf_from_storage(supabase, pdf_path)

    supabase.table("paper_tags").delete().eq("paper_id", row["id"]).execute()
    (
        supabase.table("papers")
        .delete()
        .eq("id", row["id"])
        .eq("user_id", user_id)
        .execute()
    )
